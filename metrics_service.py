"""
Flask API Wrapper for Metrics Extraction Service
Wraps extract_metrics_new.py as a REST API for Docker
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import sys
import tempfile
from werkzeug.utils import secure_filename
from extract_metrics_new import UniversalExtractorWithReport
from docx import Document

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'docx'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
AVG_TIME_PER_IMAGE = 2.5  # Average seconds per image for OCR

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def count_images_in_docx(filepath):
    """Count the number of images in a DOCX file using multiple methods"""
    try:
        doc = Document(filepath)
        image_count = 0
        
        # Method 1: Count images via document relationships (most reliable)
        # This counts actual image files embedded in the document
        image_parts = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_parts.append(rel.target_ref)
        
        relationship_count = len(image_parts)
        print(f"Method 1 (Relationships): Found {relationship_count} images")
        
        # Method 2: Count inline images in paragraphs via blip elements
        blip_count = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if 'blip' in run._element.xml:
                    blip_count += 1

        # Count images in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'blip' in run._element.xml:
                                blip_count += 1
        
        print(f"Method 2 (Blip elements): Found {blip_count} images")
        
        # Use the maximum count from both methods
        image_count = max(relationship_count, blip_count)
        
        print(f"Final image count: {image_count}")
        
        return image_count if image_count > 0 else 1  # Minimum 1 if no images found
    except Exception as e:
        print(f"Error counting images: {e}")
        import traceback
        traceback.print_exc()
        return 1


def estimate_processing_time(image_count):
    """Estimate processing time in seconds based on image count"""
    model_download_time = 30  # First run model download is ~30 seconds
    processing_time = image_count * AVG_TIME_PER_IMAGE
    return int(model_download_time + processing_time)

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'metrics-extraction'
    }), 200

@app.route('/extract', methods=['POST'])
def extract_metrics():
    """
    Extract metrics from DOCX file

    Expected: multipart/form-data with 'file' field containing DOCX file
    Returns: JSON with extracted metrics
    """
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Only .docx files are allowed'}), 400

        # Save uploaded file temporarily
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            # Create temporary output directory
            output_dir = os.path.join(app.config['UPLOAD_FOLDER'], f'extraction_{os.getpid()}')

            # Process the document
            extractor = UniversalExtractorWithReport(filepath, output_dir)
            results = extractor.process_document()
            extractor.save_results(results)

            # Format results for API response
            metrics = []
            for r in results:
                metrics.append({
                    'id': r['id'],
                    'title': r['title'],
                    'value': r['value'],
                    'image_file': r['image_file'],
                    'debug_info': r['debug_info']
                })

            return jsonify({
                'success': True,
                'metrics': metrics,
                'count': len(metrics),
                'output_dir': output_dir
            }), 200

        finally:
            # Clean up uploaded file
            if os.path.exists(filepath):
                os.remove(filepath)

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/estimate', methods=['POST'])
def estimate():
    """
    Estimate processing time based on image count
    Returns: image_count and estimated_time_seconds
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Only .docx files are allowed'}), 400

        # Save uploaded file temporarily
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            # Count images
            image_count = count_images_in_docx(filepath)
            estimated_time = estimate_processing_time(image_count)
            
            print(f"=== ESTIMATION RESPONSE ===")
            print(f"Image count: {image_count}")
            print(f"Estimated time: {estimated_time}s")
            print(f"===========================")

            return jsonify({
                'success': True,
                'image_count': image_count,
                'estimated_time_seconds': estimated_time,
                'message': f'Found {image_count} image(s). Estimated processing time: ~{estimated_time} seconds'
            }), 200

        finally:
            # Clean up uploaded file
            if os.path.exists(filepath):
                os.remove(filepath)

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/extract-simple', methods=['POST'])
def extract_simple():
    """
    Simplified extraction - returns only title and value pairs
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Only .docx files are allowed'}), 400

        # Save uploaded file temporarily
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            # Count images for metadata
            image_count = count_images_in_docx(filepath)
            estimated_time = estimate_processing_time(image_count)

            # Create temporary output directory
            output_dir = os.path.join(app.config['UPLOAD_FOLDER'], f'extraction_{os.getpid()}')

            # Process the document
            print(f"=== STARTING EXTRACTION ===")
            print(f"File: {filename}")
            print(f"Image count: {image_count}")
            
            extractor = UniversalExtractorWithReport(filepath, output_dir)
            results = extractor.process_document()
            
            print(f"Extraction complete. Results count: {len(results)}")
            
            # Debug: Print first few results
            for i, r in enumerate(results[:3]):
                print(f"Result {i+1}: title='{r['title']}', value={r['value']}, type={type(r['value'])}")

            # Return only title and value pairs (simpler format)
            metrics = [
                {'title': r['title'], 'value': r['value']}
                for r in results
            ]
            
            print(f"Formatted metrics count: {len(metrics)}")
            print(f"First metric: {metrics[0] if metrics else 'None'}")
            print(f"=== EXTRACTION COMPLETE ===")

            return jsonify({
                'success': True,
                'metrics': metrics,
                'count': len(metrics),
                'image_count': image_count,
                'estimated_time_seconds': estimated_time
            }), 200

        finally:
            # Clean up uploaded file
            if os.path.exists(filepath):
                os.remove(filepath)

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

if __name__ == '__main__':
    # Run on port 5000, accessible from Docker network
    app.run(host='0.0.0.0', port=5000, debug=False)
